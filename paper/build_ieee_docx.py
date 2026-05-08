from __future__ import annotations

from pathlib import Path
import re
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
TEX_PATH = BASE_DIR / "querymind_paper.tex"
OUT_PATH = BASE_DIR / "QueryMind_Research_Paper.docx"


def set_font(run, name: str = "Times New Roman", size: int = 10, bold: bool = False, italic: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_margins(section, left=0.7, right=0.7, top=0.75, bottom=1.0):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)


def set_two_columns(section, gap=0.25):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), str(int(Inches(gap).emu / 12700)))


def latex_to_text(text: str) -> str:
    if not text:
        return ""

    replacements = {
        r"\%": "%",
        r"\&": "&",
        r"\_": "_",
        r"\#": "#",
        r"\{": "{",
        r"\}": "}",
        r"\$": "$",
        r"\sim": "~",
        r"\textbackslash{}": "\\",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    text = re.sub(r"\\url\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\texttt\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\cite\{[^}]*\}", "", text)
    text = re.sub(r"\\ref\{[^}]*\}", "", text)
    text = re.sub(r"\\label\{[^}]*\}", "", text)
    text = text.replace("``", '"').replace("''", '"')
    text = text.replace("---", "-").replace("--", "-")
    text = text.replace("~", " ")
    text = re.sub(r"\\[a-zA-Z]+\*?(\[[^\]]*\])?(\{[^{}]*\})?", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract(source: str, start: str, end: str) -> str:
    return source.split(start, 1)[1].split(end, 1)[0].strip()


def add_paragraph(doc: Document, text: str, size: int = 10, bold: bool = False, italic: bool = False,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_after: int = 0, style: str | None = None):
    para = doc.add_paragraph(style=style)
    para.alignment = align
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    return para


def add_code_block(doc: Document, code: str):
    for line in code.strip("\n").splitlines():
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.12)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line.rstrip())
        set_font(run, name="Courier New", size=8)


def add_list_item(doc: Document, text: str, numbered: bool = False):
    style = "List Number" if numbered else "List Bullet"
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run(text)
    set_font(run, size=10)


def parse_blocks(doc: Document, text: str):
    text = text.strip()
    if not text:
        return

    parts = re.split(r"(\\begin\{lstlisting\}.*?\\end\{lstlisting\})", text, flags=re.S)
    for part in parts:
        if not part.strip():
            continue
        if part.startswith(r"\\begin{lstlisting}"):
            match = re.search(r"\\begin\{lstlisting\}(.*?)\\end\{lstlisting\}", part, flags=re.S)
            if match:
                add_code_block(doc, match.group(1))
            continue

        cleaned = latex_to_text(part)
        if not cleaned:
            continue

        paragraphs = [chunk.strip() for chunk in re.split(r"\n\s*\n", cleaned) if chunk.strip()]
        for para_text in paragraphs:
            lines = [line.strip() for line in para_text.splitlines() if line.strip()]
            if not lines:
                continue

            if all(line.startswith("-") for line in lines):
                for line in lines:
                    add_list_item(doc, line.lstrip("- "))
                continue

            if all(re.match(r"^\d+\.", line) for line in lines):
                for line in lines:
                    add_list_item(doc, re.sub(r"^\d+\.\s*", "", line), numbered=True)
                continue

            add_paragraph(doc, " ".join(lines), size=10, space_after=2)


def main():
    source = TEX_PATH.read_text(encoding="utf-8")
    title = latex_to_text(source.split(r"\title{", 1)[1].split(r"\author{", 1)[0]).replace(r"\\", "\n").rstrip("}")
    abstract = extract(source, r"\begin{abstract}", r"\end{abstract}")
    keywords = extract(source, r"\begin{IEEEkeywords}", r"\end{IEEEkeywords}")
    body = source.split(r"\end{IEEEkeywords}", 1)[1].split(r"\begin{thebibliography}", 1)[0]
    bibliography = extract(source, r"\begin{thebibliography}{00}", r"\end{thebibliography}")

    doc = Document()
    section = doc.sections[0]
    set_margins(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10)

    add_paragraph(doc, title, size=24, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for idx, line in enumerate(["[Author Name]", "[Institution Name]", "[email@institution.edu]"]):
        run = author_para.add_run(line)
        set_font(run, size=10)
        if idx < 2:
            run.add_break(WD_BREAK.LINE)
    author_para.paragraph_format.space_after = Pt(8)

    add_paragraph(doc, "Abstract", size=9, bold=True, space_after=2)
    add_paragraph(doc, latex_to_text(abstract), size=10, space_after=5)
    add_paragraph(doc, "Index Terms", size=9, bold=True, space_after=2)
    add_paragraph(doc, latex_to_text(keywords), size=10, space_after=8)

    doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_margins(doc.sections[-1])
    set_two_columns(doc.sections[-1], gap=0.25)

    section_blocks = re.split(r"(?=\\section\*?\{)", body)
    for block in section_blocks:
        block = block.strip()
        if not block:
            continue
        match = re.match(r"\\section\*?\{([^}]*)\}", block)
        if not match:
            continue

        add_paragraph(doc, latex_to_text(match.group(1)), size=10, bold=True, space_after=4)
        content = block[match.end():]
        subsections = list(re.finditer(r"\\subsection\*?\{([^}]*)\}", content))

        if not subsections:
            parse_blocks(doc, content)
            continue

        parse_blocks(doc, content[:subsections[0].start()])
        for idx, sub in enumerate(subsections):
            sub_title = latex_to_text(sub.group(1))
            sub_start = sub.end()
            sub_end = subsections[idx + 1].start() if idx + 1 < len(subsections) else len(content)
            add_paragraph(doc, sub_title, size=10, bold=True, italic=True, space_after=2)
            parse_blocks(doc, content[sub_start:sub_end])

    add_paragraph(doc, "References", size=10, bold=True, space_after=4)
    refs = re.findall(r"\\bibitem\{[^}]*\}\s*(.*?)(?=\\bibitem\{|$)", bibliography, flags=re.S)
    for idx, ref in enumerate(refs, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.18)
        para.paragraph_format.first_line_indent = Inches(-0.18)
        para.paragraph_format.space_after = Pt(0)
        ref_text = latex_to_text(ref).replace("\n", " ")
        run = para.add_run(f"[{idx}] {ref_text}")
        set_font(run, size=9)

    doc.save(OUT_PATH)
    print(f"Created {OUT_PATH}")


if __name__ == "__main__":
    main()